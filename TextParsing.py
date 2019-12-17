from docx import Document
import re
import json
from pyresparser import ResumeParser


def text2word(path_to_text):
    document = Document()
    # document.add_heading(i, 0)
    textfile = open(path_to_text).read()
    textfile = re.sub(r'[^\x00-\x7F]+|\x0c', ' ', textfile)
    p = document.add_paragraph(textfile)
    path_to_docx = path_to_text.replace(".txt", ".docx")
    document.save(path_to_docx)
    return path_to_docx


def resume_parser(path_to_docx):
    data = ResumeParser(path_to_docx).get_extracted_data()
    path_to_json = path_to_docx.replace(".docx", ".json")
    with open(path_to_json, 'w') as json_file:
        json.dump(data, json_file)
    return data


def text_parsing(path_to_text):
    # convert text to docx to use pyresparser
    path_to_docx = text2word(path_to_text)

    # resume parser
    data = resume_parser(path_to_docx)

    return data
